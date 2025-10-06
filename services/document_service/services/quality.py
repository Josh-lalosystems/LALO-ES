"""
Document quality scoring utilities.

This module provides a lightweight heuristic quality scorer for incoming
documents. It attempts inexpensive checks (extract text length, OCR heuristics
like ratio of printable chars, presence of many images, page counts) and
returns a score and category to route documents into different processing
pipelines (full, light, manual_review).

The scorer is intentionally conservative and safe: failures return a low but
non-zero score and never raise.
"""
from typing import Tuple, Dict
import re
import io
import logging

logger = logging.getLogger('doc_quality')


def _printable_ratio(s: str) -> float:
    if not s:
        return 0.0
    printable = re.sub(r"\s+", "", s)
    if len(printable) == 0:
        return 0.0
    # ratio of letters/digits/punct to total
    good = sum(1 for ch in printable if ch.isprintable())
    return good / len(printable)


def score_document(content_bytes: bytes, mime_type: str = "") -> Tuple[int, Dict]:
    """Score a document and return (score, details).

    Score 0-100 where higher is better. Details contains heuristic metrics.
    """
    try:
        text_estimate = ""
        page_count = None
        has_images = False
        tables = 0

        # Lightweight PDF handling
        if "pdf" in mime_type.lower():
            try:
                import PyPDF2
                from io import BytesIO
                reader = PyPDF2.PdfReader(BytesIO(content_bytes))
                page_count = len(reader.pages)
                # quick text concat from first few pages to estimate quality
                sample_pages = min(5, page_count)
                texts = []
                for i in range(sample_pages):
                    try:
                        texts.append(reader.pages[i].extract_text() or "")
                    except Exception:
                        texts.append("")
                text_estimate = "\n".join(texts)
                # detect images heuristic
                try:
                    has_images = any(len(p.images) > 0 for p in reader.pages[:sample_pages])
                except Exception:
                    has_images = False
            except Exception:
                text_estimate = ""

        # Word (docx)
        elif "word" in mime_type.lower() or "docx" in mime_type.lower():
            try:
                from docx import Document as DocxDocument
                from io import BytesIO
                doc = DocxDocument(BytesIO(content_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text]
                text_estimate = "\n".join(paragraphs[:40])
                page_count = len(doc.paragraphs)
                # tables count
                try:
                    tables = len(doc.tables)
                except Exception:
                    tables = 0
            except Exception:
                text_estimate = ""

        # Excel
        elif "spreadsheet" in mime_type.lower() or "excel" in mime_type.lower():
            try:
                import openpyxl
                from io import BytesIO
                wb = openpyxl.load_workbook(BytesIO(content_bytes), read_only=True)
                # Collect limited cell text to estimate quality
                snippets = []
                for i, sheet in enumerate(wb.sheetnames[:3]):
                    ws = wb[sheet]
                    for row in ws.iter_rows(min_row=1, max_row=20, max_col=8, values_only=True):
                        row_text = " ".join([str(c) for c in row if c is not None])
                        if row_text:
                            snippets.append(row_text)
                text_estimate = "\n".join(snippets[:100])
                page_count = len(wb.sheetnames)
                tables = 1 if snippets else 0
            except Exception:
                text_estimate = ""

        # Fallback: try to decode as utf-8 text
        else:
            try:
                text_estimate = content_bytes.decode('utf-8', errors='ignore')[:2000]
            except Exception:
                text_estimate = ""

        words = len(text_estimate.split()) if text_estimate else 0
        printable_ratio = _printable_ratio(text_estimate)

        # Heuristic score composition
        score = 0
        # Base on text quantity
        if words > 2000:
            score += 50
        elif words > 500:
            score += 30
        elif words > 100:
            score += 10

        # Printable ratio
        if printable_ratio > 0.9:
            score += 30
        elif printable_ratio > 0.7:
            score += 15
        elif printable_ratio > 0.4:
            score += 5

        # Penalize enormous image-only PDFs
        if page_count and page_count > 0 and not text_estimate and has_images:
            score = min(score, 20)

        # Tables bump score if present and some text exists
        if tables and words > 10:
            score += 10

        # Normalize
        score = max(0, min(100, score))

        # Category mapping
        if score >= 60:
            category = "clean"
        elif score >= 30:
            category = "decent"
        else:
            category = "garbage"

        details = {
            "words": words,
            "printable_ratio": printable_ratio,
            "pages": page_count,
            "has_images": bool(has_images),
            "tables": int(tables),
            "category": category
        }

        return int(score), details

    except Exception as e:
        logger.warning("Quality scoring failed: %s", e)
        return 20, {"error": str(e)}
