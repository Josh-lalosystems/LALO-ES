"""
Copyright (c) 2025 LALO AI SYSTEMS, LLC. All rights reserved.

PROPRIETARY AND CONFIDENTIAL

This file is part of LALO AI Platform and is protected by copyright law.
Unauthorized copying, modification, distribution, or use of this software,
via any medium, is strictly prohibited without the express written permission
of LALO AI SYSTEMS, LLC.
"""

from typing import Dict, Any
from io import BytesIO
from docx import Document as DocxDocument
from ...models import Document, ProcessingResult
from ..chunker import chunk_text_hierarchical
from ..table_extractor import word_table_to_chunks

class WordProcessor:
    async def process(self, document: Document) -> ProcessingResult:
        try:
            # Load document from bytes
            doc = DocxDocument(BytesIO(document.content))
            
            # Extract metadata
            metadata: Dict[str, Any] = {
                "paragraphs": len(doc.paragraphs),
                "sections": len(doc.sections),
                "tables": len(doc.tables),
                "styles": [],
                "properties": {}
            }
            
            # Get styles
            for style in doc.styles:
                if not style.name.startswith('@'):
                    metadata["styles"].append(style.name)
            
            # Get document properties
            if doc.core_properties:
                metadata["properties"] = {
                    "title": doc.core_properties.title,
                    "author": doc.core_properties.author,
                    "subject": doc.core_properties.subject,
                    "keywords": doc.core_properties.keywords,
                    "category": doc.core_properties.category,
                    "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else None
                }
            
            # Update document metadata
            document.metadata.update(metadata)

            # Build text and chunk it
            try:
                paragraphs = [p.text for p in doc.paragraphs if p.text]
                full_text = "\n\n".join(paragraphs)
                chunks = chunk_text_hierarchical(full_text, doc_id=document.id)
                # Also extract tables into chunks
                try:
                    table_chunks = word_table_to_chunks(doc.tables, doc_id=document.id)
                    document.metadata['chunks'] = chunks + table_chunks
                except Exception:
                    document.metadata['chunks'] = chunks
            except Exception:
                document.metadata['chunks'] = []
            
            return ProcessingResult(
                success=True,
                document_id=document.id,
                message="Word document processed successfully"
            )
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                document_id=document.id,
                errors=[f"Word processing error: {str(e)}"]
            )
