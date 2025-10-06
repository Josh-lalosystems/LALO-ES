import io
from docx import Document as DocxDocument
import openpyxl

from services.document_service.services.table_extractor import word_table_to_chunks, excel_tables_to_chunks
from services.document_service.services.chunker import chunk_text_hierarchical


def test_word_table_to_chunks_and_deterministic_ids():
    doc = DocxDocument()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"

    chunks = word_table_to_chunks(doc.tables, doc_id="testdoc")
    assert len(chunks) == 1
    c = chunks[0]
    assert c['level'] == 'table'
    assert 'chunk_id' in c and isinstance(c['chunk_id'], str)
    # deterministic: calling again should produce same id
    chunks2 = word_table_to_chunks(doc.tables, doc_id="testdoc")
    assert chunks2[0]['chunk_id'] == c['chunk_id']
    assert c['token_count'] > 0


def test_excel_tables_to_chunks_and_token_count(tmp_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Col1", "Col2"])
    ws.append(["X", "Y"])

    chunks = excel_tables_to_chunks(wb, doc_id="exceldoc")
    assert len(chunks) >= 1
    c = chunks[0]
    assert c['level'] == 'table'
    assert c['token_count'] > 0
    # deterministic id check
    chunks2 = excel_tables_to_chunks(wb, doc_id="exceldoc")
    assert chunks2[0]['chunk_id'] == c['chunk_id']
